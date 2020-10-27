#!/usr/bin/env python
# -*-coding:utf-8-*-
"""
status = '0'    正常
status = '10'   逻辑删除
status = '20'   禁用
"""

import json
import shortuuid
import base64
from websdk.jwt_token import gen_md5
from websdk.tools import check_password
from libs.base_handler import BaseHandler
from websdk.db_context import DBContext
# from models.admin import Users, UserRoles,model_to_dict as users_model_to_dict
from models.server import AssetSql, model_to_dict
from models.db import DB, model_to_dict    as   DB_model_to_dict
from websdk.consts import const
from websdk.cache_context import cache_conn
from websdk.web_logs import ins_log
from sqlalchemy import or_, and_
from sqlalchemy import func
import tornado.ioloop
import tornado.web
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import shared
import os


class SqlListHandler(BaseHandler):
    def post(self, *args, **kwargs):
        data = json.loads(self.request.body.decode("utf-8"))
        name = data.get('name', None)  # 个案名称
        header = data.get('header', None)
        dbname_id = data.get('dbname_id', None)
        dbname = data.get('dbname', None)
        totype = data.get('totype', None)
        sqlstr = data.get('sqlstr', None)  # 优先级
        remarks = data.get('remarks', False)  # 执行人
        username = data.get('username', False)  # 执行人
        obj = data.get('obj', False)  # 项目
        department = data.get('department', False)  # 部门
        storage = data.get('storage', False)  # 部门
        create_time = data.get('create_time', False)
        # ins_log.read_log('info', "800000000000000000000000000000000000")
        with DBContext('w', None, True) as session:
            session.add(AssetSql(
                name=name,
                header=header,
                dbname_id=dbname_id,
                dbname=dbname,
                totype=totype,
                sqlstr=sqlstr,
                remarks=remarks,
                username=username,
                obj=obj,
                department=department,
                storage=storage,
                create_time=create_time,))
            session.commit()
        self.write(dict(code=0, msg='成功', count=0, data=[]))

    def put(self, *args, **kwargs):
        data = json.loads(self.request.body.decode("utf-8"))
        id = int(data.get('id', None))  # id号
        name = data.get('name', None)
        header = data.get('header', None)
        dbname_id = data.get('dbname_id', None)
        dbname = data.get('dbname', None)
        totype = data.get('totype', None)
        sqlstr = data.get('sqlstr', None)
        remarks = data.get('remarks', False)
        username= data.get('username', False)
        obj = data.get('obj', False)
        department = data.get('department', False)
        storage = data.get('storage', False)
        create_time = data.get('create_time', False)
        with DBContext('w', None, True) as session:
            session.query(AssetSql).filter(AssetSql.id == id).update({
                AssetSql.name: name,
                AssetSql.header: header,
                AssetSql.dbname_id: dbname_id,
                AssetSql.dbname: dbname,
                AssetSql.totype: totype,
                AssetSql.sqlstr: sqlstr,
                AssetSql.remarks: remarks,
                AssetSql.username: username,
                AssetSql.obj: obj,
                AssetSql.department: department,
                AssetSql.storage: storage,
                AssetSql.create_time: create_time,
            })
            session.commit()
        self.write(dict(code=0, msg='成功', count=0, data=[]))


class getSqlListHandler(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        superuser_flag  = 0
        tovalue = self.get_argument('value', strip=True)  # 要查询的关键字
        topage = int(self.get_argument('page', strip=1))  # 开始页
        tolimit = int(self.get_argument('limit', strip=10))  # 要查询条数
        limit_start = (topage - 1) * tolimit

        with DBContext('r') as session:
            params = {}
            if tovalue:
                params = eval(tovalue)
            conditions = []
            if self.is_superuser:
                superuser_flag = 1
            if params.get('name', ''):
                conditions.append(AssetSql.name.like('%{}%'.format(params['name'])))
            if params.get('remarks', ''):
                conditions.append(AssetSql.remarks.like('%{}%'.format(params['remarks'])))
            if params.get('totype', ''):
                conditions.append(AssetSql.totype.like('%{}%'.format(params['totype'])))

            todata = session.query(AssetSql).filter(*conditions).order_by(AssetSql.create_time.desc()).offset(limit_start).limit(int(tolimit)).all()
            tocount = session.query(AssetSql).filter(*conditions).count()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["name"] = data_dict["name"]
            case_dict["header"] = data_dict["header"]
            case_dict["dbname_id"] = data_dict["dbname_id"]
            case_dict["dbname"] = data_dict["dbname"]
            case_dict["totype"] = data_dict["totype"]
            case_dict["sqlstr"] = data_dict["sqlstr"]
            case_dict["remarks"] = data_dict["remarks"]
            case_dict["username"] = data_dict["username"]
            case_dict["obj"] = data_dict["obj"]
            case_dict["department"] = data_dict["department"]
            case_dict["storage"] = data_dict["storage"]
            case_dict["create_time"] = str(data_dict["create_time"])
            data_list.append(case_dict)

        if len(data_list) > 0:
            self.write(dict(code=0, msg='获取成功', count=tocount, data=data_list,flag=superuser_flag))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[],flag=superuser_flag))


class getCasefileHandler(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        nickname = self.get_current_nickname()
        tostart = self.get_argument('startdate', strip=True)  # 要查询的关键字
        toend = self.get_argument('enddate', strip=True)  # 要查询的关键字
        # CaseList表
        with DBContext('r') as session:
            params = {}
            conditions = []
            # if self.is_superuser:
            #     pass
            # else:
            #     conditions.append(AssetSql.case_executor == nickname)
            conditions.append(AssetSql.case_stime >= tostart)
            conditions.append(AssetSql.case_etime <= toend)
            todata = session.query(AssetSql).filter(*conditions).order_by(AssetSql.ctime.desc()).all()
        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["case_num"] = data_dict["case_num"]
            case_dict["case_obj"] = data_dict["case_obj"]
            case_dict["demand_unit"] = data_dict["demand_unit"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_type"] = data_dict["case_type"]
            case_dict["case_ltime"] = data_dict["case_ltime"]
            case_dict["case_name"] = data_dict["case_name"]
            case_dict["case_status"] = data_dict["case_status"]
            case_dict["case_priority"] = data_dict["case_priority"]
            case_dict["demander"] = data_dict["demander"]
            case_dict["case_executor"] = data_dict["case_executor"]
            case_dict["case_source"] = data_dict["case_source"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_stime"] = str(data_dict["case_stime"])
            case_dict["case_etime"] = str(data_dict["case_etime"])
            case_dict["case_creator"] = data_dict["case_creator"]
            data_list.append(case_dict)
        # PlanList表
        plandata_list = []  # 计划工作完成情况
        with DBContext('r') as session:
            params = {}
            conditions = []
            # if self.is_superuser:
            #     pass
            # else:
            #     conditions.append(AssetSql.case_executor == nickname)
            conditions.append(PlanList.plan_stime >= tostart)
            conditions.append(PlanList.plan_etime <= toend)
            conditions.append(PlanList.plan_status == "处理中")
            todata = session.query(PlanList).filter(*conditions).order_by(PlanList.ctime.desc()).all()
        for msg in todata:
            plan_dict = {}
            data_dict = model_to_dict(msg)
            plan_dict["id"] = data_dict["id"]
            plan_dict["plan_num"] = data_dict["plan_num"]
            plan_dict["plan_obj"] = data_dict["plan_obj"]
            plan_dict["demand_unit"] = data_dict["demand_unit"]
            plan_dict["plan_details"] = data_dict["plan_details"]
            plan_dict["plan_type"] = data_dict["plan_type"]
            plan_dict["plan_ltime"] = data_dict["plan_ltime"]
            plan_dict["plan_name"] = data_dict["plan_name"]
            plan_dict["plan_status"] = data_dict["plan_status"]
            plan_dict["plan_priority"] = data_dict["plan_priority"]
            plan_dict["demander"] = data_dict["demander"]
            plan_dict["plan_executor"] = data_dict["plan_executor"]
            plan_dict["plan_source"] = data_dict["plan_source"]
            plan_dict["plan_details"] = data_dict["plan_details"]
            plan_dict["plan_stime"] = str(data_dict["plan_stime"])
            plan_dict["plan_etime"] = str(data_dict["plan_etime"])
            plan_dict["plan_creator"] = data_dict["plan_creator"]
            plandata_list.append(plan_dict)
        ins_log.read_log('info', plandata_list)
        if (len(data_list) + len(plandata_list)) > 0:
            import docx
            import time, datetime
            flag = 0  # 周报0月报1
            # ins_log.read_log('info', "800000000000000000000000000000000000")
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/template/运维组工作报告模板.docx'.format(Base_DIR)
            # ins_log.read_log('info', upload_path)
            # ins_log.read_log('info', "8000000000011111111111111111111")
            # doc = docx.Document(u"/opt/codo/codo-problem/static/report/template/运维组工作报告模板.docx")
            doc = docx.Document(upload_path)
            if (int(str(toend)[8:10]) - int(str(tostart)[8:10])) <= 7:
                doc.paragraphs[0].text = ['工作周报']
                flag = 0
            else:
                doc.paragraphs[0].text = ['工作月报']
                flag = 1
            tempstr = '报告时间：' + str(tostart)[0:4] + '年' + str(tostart)[5:7] + '月' + str(tostart)[8:10] + '日' \
                      + '～' + str(toend)[0:4] + '年' + str(toend)[5:7] + '月' + str(toend)[8:10] + '日'
            # ins_log.read_log('info', tempstr)
            tempdict = []
            tempdict.append(tempstr)
            # ins_log.read_log('info', tempdict)
            doc.paragraphs[3].text = tempdict
            tempstr = ''
            tempstr = '报 告 人：' + nickname
            tempdict = []
            tempdict.append(tempstr)
            doc.paragraphs[4].text = tempdict
            # 本周/月工作完成情况（2020-05-21~2020-05-31）
            tempstr = ''
            if flag == 0:
                tempstr = '本周工作完成情况（'
            else:
                tempstr = '本月工作完成情况（'

            tempstr = tempstr + str(tostart)[0:4] + '-' + str(tostart)[5:7] + '-' + str(tostart)[8:10] \
                      + '～' + str(toend)[0:4] + '-' + str(toend)[5:7] + '-' + str(toend)[8:10] + '）'
            tempdict = []
            tempdict.append(tempstr)
            doc.paragraphs[6].text = tempdict
            # 下周/月工作计划（2020-xx-xx~2020-xx-xx）
            # ins_log.read_log('info', toend)
            if flag == 0:
                tempstr = ''
                tempstr = toend + ' ' + '23:59:59'
                # ins_log.read_log('info', tempstr)
                timeArray = time.strptime(tempstr, "%Y-%m-%d %H:%M:%S")
                # 转为时间戳
                timeStamp = int(time.mktime(timeArray))
                # ins_log.read_log('info', timeStamp)
                # 下周一
                timeStamp1 = timeStamp + 3600 * 24
                startlocaltime = time.localtime(timeStamp1)
                startdatatime = time.strftime("%Y-%m-%d %H:%M:%S", startlocaltime)
                # ins_log.read_log('info', startdatatime)
                # 下周日
                timeStamp2 = timeStamp + 3600 * 24 * 7
                endlocaltime = time.localtime(timeStamp2)
                enddatatime = time.strftime("%Y-%m-%d %H:%M:%S", endlocaltime)
                # ins_log.read_log('info', enddatatime)
                tempstr = ''
                tempstr = '下周工作计划（' + str(startdatatime)[0:4] + '-' + str(startdatatime)[5:7] + '-' + str(
                    startdatatime)[8:10] \
                          + '～' + str(enddatatime)[0:4] + '-' + str(enddatatime)[5:7] + '-' + str(enddatatime)[
                                                                                              8:10] + '）'
            else:
                import calendar
                toyear = int(str(toend)[0:4])
                tomonth = int(str(toend)[5:7]) + 1
                if tomonth > 12:
                    tomonth = 1
                    toyear = toyear + 1
                monthRange = calendar.monthrange(toyear, tomonth)  # 返回值是元组，第一个参数是这个月的第一天是星期几，第二个参数是这个月的总天数

                tempstr = ''
                if len(str(tomonth)) == 1:
                    tomonthstr = '0' + str(tomonth)
                else:
                    tomonthstr = str(tomonth)
                tempstr = '下个月工作计划（' + str(toyear) + '-' + tomonthstr + '-' + "01" \
                          + '～' + str(toyear) + '-' + tomonthstr + '-' + str(monthRange[1]) + '）'
            tempdict = []
            tempdict.append(tempstr)
            doc.paragraphs[7].text = tempdict

            rows_index = 0  # 行数
            merge_index = 0  # 合并个数
            totable = doc.tables[0]
            laiwen_list = []  # 来问
            upgrade_list = []  # 升级
            fault_list = []  # 故障

            especially_list = []  # 特急/耗时长的情况
            sudden_list = []  # 其他突发工作
            # data_list.extend(plan_list) #合并两个列表
            for i in data_list:
                if i["case_source"] == '来文':
                    laiwen_list.append(i)
                    totable.add_row()
                if i["case_type"] == '应用升级':
                    upgrade_list.append(i)
                    totable.add_row()
                if i["case_type"] == '故障':
                    fault_list.append(i)
                    totable.add_row()
                if i["case_priority"] == '特急' or int(i["case_ltime"]) >= 240:
                    especially_list.append(i)
                    totable.add_row()
            for i in plandata_list:
                totable.add_row()

            rows_index += 1  # 行数
            merge_index += len(laiwen_list)  # 合并个数
            if len(laiwen_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "来文情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for k in range(0, len(laiwen_list)):
                    totable.cell(rows_index, 1).text = "来文情况"
                    totable.cell(k + 1, 2).text = str(k + 1) + '.' + laiwen_list[k]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(len(laiwen_list), 1))

            rows_index += len(laiwen_list)  # 行数
            merge_index += len(upgrade_list)  # 合并个数
            if len(upgrade_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "升级情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for j in range(0, len(upgrade_list)):
                    totable.cell(rows_index, 1).text = "升级情况"
                    totable.cell(rows_index + j, 2).text = str(j + 1) + '.' + upgrade_list[j]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(upgrade_list)  # 行数
            merge_index += len(fault_list)  # 合并个数
            if len(fault_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "故障情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for h in range(0, len(fault_list)):
                    totable.cell(rows_index, 1).text = "故障情况"
                    totable.cell(rows_index + h, 2).text = str(h + 1) + '.' + fault_list[h]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(fault_list)  # 行数
            merge_index += len(plandata_list)  # 合并个数
            if len(plandata_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "计划工作完成情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for g in range(0, len(plandata_list)):
                    totable.cell(rows_index, 1).text = "计划工作完成情况"
                    totable.cell(rows_index + g, 2).text = str(g + 1) + '.' + plandata_list[g]["plan_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(plandata_list)  # 行数
            merge_index += len(especially_list)  # 合并个数
            if len(especially_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "重要工作情况"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for g in range(0, len(especially_list)):
                    totable.cell(rows_index, 1).text = "重要工作情况"
                    totable.cell(rows_index + g, 2).text = str(g + 1) + '.' + especially_list[g]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            rows_index += len(especially_list)  # 行数
            merge_index += len(sudden_list)  # 合并个数
            if len(sudden_list) == 0:
                totable.add_row()
                totable.cell(rows_index, 1).text = "其他突发工作"
                rows_index += 1  # 行数
                merge_index += 1  # 合并个数
            else:
                for g in range(0, len(sudden_list)):
                    totable.cell(rows_index, 1).text = "其他突发工作"
                    totable.cell(rows_index + g, 2).text = str(g + 1) + '.' + sudden_list[g]["case_name"]
                totable.cell(rows_index, 1).merge(totable.cell(merge_index, 1))

            # ins_log.read_log('info', especially_list)
            # ins_log.read_log('info', len(especially_list))
            totable.cell(1, 0).merge(totable.cell(merge_index, 0))

            tempstr = ''
            tempstr = '维护组工作报告_' + nickname + '[' + str(tostart)[0:4] + str(tostart)[5:7] + str(tostart)[8:10] \
                      + '-' + str(toend)[0:4] + str(toend)[5:7] + str(toend)[8:10] + ']' + '.docx'
            # ins_log.read_log('info', tempstr)
            # 个案数据统计分析
            with DBContext('r') as session:
                conditions = []
                conditions.append(AssetSql.case_source == "来文")
                conditions.append(AssetSql.case_stime >= tostart)
                conditions.append(AssetSql.case_etime <= toend)
                lime = session.query(func.count(AssetSql.case_source),
                                     func.date_format(AssetSql.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()
            with DBContext('r') as session:
                conditions = []
                conditions.append(AssetSql.case_type == "应用升级")
                conditions.append(AssetSql.case_stime >= tostart)
                conditions.append(AssetSql.case_etime <= toend)
                lime2 = session.query(func.count(AssetSql.case_type),
                                      func.date_format(AssetSql.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()
            with DBContext('r') as session:
                conditions = []
                conditions.append(AssetSql.case_type == "故障")
                conditions.append(AssetSql.case_stime >= tostart)
                conditions.append(AssetSql.case_etime <= toend)
                lime3 = session.query(func.count(AssetSql.case_type),
                                      func.date_format(AssetSql.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()
            with DBContext('r') as session:
                conditions = []
                conditions.append(or_(AssetSql.case_priority == "特急", AssetSql.case_ltime >= 240))
                conditions.append(AssetSql.case_stime >= tostart)
                conditions.append(AssetSql.case_etime <= toend)
                lime4 = session.query(func.count(AssetSql.case_priority),
                                      func.date_format(AssetSql.case_stime, '%Y-%m-%d').label('date')).filter(
                    *conditions).group_by('date').all()

            # ins_log.read_log('info', lime)
            # ins_log.read_log('info', lime2)
            # 生成开始到结束的日期列表
            begin_date = str(tostart).replace('-', "")
            end_date = str(toend).replace('-', "")
            date_list = []
            begin_date = datetime.datetime.strptime(begin_date, "%Y%m%d")
            end_date = datetime.datetime.strptime(end_date, "%Y%m%d")
            while begin_date <= end_date:
                date_str = begin_date.strftime("%Y-%m-%d")
                date_list.append(date_str)
                begin_date += datetime.timedelta(days=1)
            ins_log.read_log('info', date_list)

            x_lime = date_list
            y_lime = getlist(date_list, lime)
            y_lime2 = getlist(date_list, lime2)
            y_lime3 = getlist(date_list, lime3)
            y_lime4 = getlist(date_list, lime4)
            ins_log.read_log('info', y_lime)
            ins_log.read_log('info', x_lime)
            ins_log.read_log('info', y_lime2)
            # 解决中文显示问题
            from matplotlib.font_manager import FontProperties
            font = FontProperties(fname=r"/usr/share/fonts/china.ttf", size=12)

            plt.figure(figsize=(7, 11))
            plt.xlabel("日期", fontproperties=font)
            plt.ylabel("数量", fontproperties=font)
            plt.plot(x_lime, y_lime, color='r', label="来文", )
            plt.plot(x_lime, y_lime2, color='cyan', label="升级", )
            plt.plot(x_lime, y_lime3, color='b', label="故障", )
            plt.plot(x_lime, y_lime4, color='lime', label="重要工作", )
            plt.title("个案各类情况分布图", fontproperties=font)
            plt.legend(prop=font)  # 添加图例
            plt.xticks(range(0, len(x_lime)), x_lime, rotation=45)

            import time
            time_str = time.strftime('%Y.%m.%d-%H:%M:%S', time.localtime(time.time()))
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/imges/report_form'.format(Base_DIR)
            # filename = "/opt/codo/codo-problem/static/report/files/imges/report_form" + str(time_str)  + ".png"
            filename = upload_path + str(time_str) + ".png"
            plt.savefig(filename)
            # plt.show()
            paragraph = doc.paragraphs[7]  # 获取段落位置
            # p = paragraph.insert_paragraph_before(text="个案情况:每周/每月分布图")
            p = paragraph.insert_paragraph_before(text=" ")
            # p.runs[0].bold = True
            paragraph = doc.paragraphs[8]  # 获取段落位置
            p = paragraph.insert_paragraph_before(text="   ")

            # paragraph = doc.paragraphs[8]  # 获取段落位置
            # paragraph.runs[-1].add_picture(filename, width=shared.Inches(6))  # 在runs的最后一段文字后添加图片/

            # 饼状图
            sum_all = len(data_list)  # 总数
            laiwen = len(laiwen_list) / sum_all  # 来文比例
            upgrade = len(upgrade_list) / sum_all  # 升级比例
            fault = len(fault_list) / sum_all  # 故障比例
            especially = len(especially_list) / sum_all  # 特急比例
            tother = 1 - laiwen - upgrade - fault - especially
            font = FontProperties(fname=r"/usr/share/fonts/china.ttf", size=12)
            label = u'来文', u'升级', u'故障', u'重要工作', u'其他'  # 各类别标签
            color = 'red', 'orange', 'yellow', 'green', 'blue'  # 各类别颜色
            size = [laiwen, upgrade, fault, especially, tother]  # 各类别占比
            explode = (0, 0, 0, 0, 0.2)  # 各类别的偏移半径
            # 绘制饼状图
            plt.figure()  # 清空画板
            plt.figure(figsize=(8, 10))
            patches, l_text, p_text = plt.pie(size, colors=color, explode=explode, labels=label, shadow=True,
                                              autopct='%1.1f%%', startangle=90)
            for t in l_text:
                t.set_fontproperties(font)  # 把每个文本设成中文字体
            # plt.axis('equal')
            plt.title(u'工作完成情况', fontproperties=font, fontsize=12)
            plt.legend(prop=font, loc='lower left', fontsize=12)  # 图例
            # 设置legend的字体大小
            leg = plt.gca().get_legend()
            ltext = leg.get_texts()
            plt.setp(ltext, fontsize=6)

            time_str = time.strftime('%Y.%m.%d-%H:%M:%S', time.localtime(time.time()))
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/imges/report_pie'.format(Base_DIR)
            # filename = "/opt/codo/codo-problem/static/report/files/imges/report_pie" + str(time_str) + ".png"
            filename = upload_path + str(time_str) + ".png"
            plt.savefig(filename)
            # plt.show()

            paragraph = doc.paragraphs[9]  # 获取段落位置
            p = paragraph.insert_paragraph_before(
                text="比例饼图:本月/周完成来文登记" + str(len(laiwen_list)) + "次,平台升级" + str(len(upgrade_list)) + "次，同时处理故障" + str(
                    len(fault_list)) + "条,并汇总处理情况。另，也按领导要求完成其他重要工作" + str(len(especially_list)) + "项。")
            p.runs[0].bold = True
            paragraph = doc.paragraphs[10]  # 获取段落位置
            p = paragraph.insert_paragraph_before(text="  ")
            paragraph = doc.paragraphs[10]  # 获取段落位置
            paragraph.runs[-1].add_picture(filename, width=shared.Inches(6))  # 在runs的最后一段文字后添加图片/
            # 柱状图
            label = ['来文', '升级', '故障', '重要工作', '其他']  # 各类别标签
            plt.figure()  # 清空画板
            plt.figure(figsize=(6, 9))
            sum_other = sum_all - len(laiwen_list) - len(upgrade_list) - len(fault_list) - len(especially_list)
            data = [len(laiwen_list), len(upgrade_list), len(fault_list), len(especially_list), sum_other]
            plt.bar(range(len(data)), data, color=['r', 'g', 'b', 'gold'])  # or `color=['r', 'g', 'b']`

            plt.xticks(range(0, len(label)), label, fontproperties=font, fontsize=12)
            # plt.xlabel("个案情况", fontproperties=font,fontsize=12)
            plt.ylabel("数量", fontproperties=font, fontsize=12)
            plt.title("工作完成情况", fontproperties=font, fontsize=12)
            time_str = time.strftime('%Y.%m.%d-%H:%M:%S', time.localtime(time.time()))
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/imges/report_bar'.format(Base_DIR)
            # filename = "/opt/codo/codo-problem/static/report/files/imges/report_bar" + str(time_str) + ".png"
            filename = upload_path + str(time_str) + ".png"
            plt.savefig(filename)

            paragraph = doc.paragraphs[11]  # 获取段落位置
            p = paragraph.insert_paragraph_before(
                text="个案柱状图:本月/周完成来文登记" + str(len(laiwen_list)) + "次,平台升级" + str(len(upgrade_list)) + "次，同时处理故障" + str(
                    len(fault_list)) + "条,并汇总处理情况。另，也按领导要求完成其他重要工作" + str(len(especially_list)) + "项。")
            p.runs[0].bold = True
            paragraph = doc.paragraphs[12]  # 取段落位置 获
            p = paragraph.insert_paragraph_before(text="  ")
            paragraph = doc.paragraphs[12]  # 获取段落位置
            paragraph.runs[-1].add_picture(filename, width=shared.Inches(5))  # 在runs的最后一段文字后添加图片/
            Base_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            upload_path = '{}/static/report/files/'.format(Base_DIR)
            # doc.save(u"/opt/codo/codo-problem/static/report/files/" + tempstr)  # 保存文档
            doc.save(upload_path + tempstr)  # 保存文档
            # http://192.168.2.200:8200/static/report/files/%E7%BB%B4%E6%8A%A4%E7%BB%84%E5%B7%A5%E4%BD%9C%E6%8A%A5%E5%91%8A_admin[20200601-20200630].docx
            # ins_log.read_log('info',self.request.host)
            urlstr = "http://" + self.request.host + "/static/report/files/" + tempstr
            self.write(dict(code=0, msg='获取報告成功', count=1, data=urlstr))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))


def getlist(date_list, lime):
    sum_list = []
    sum_flag = 0
    for d in date_list:
        for k in lime:
            if d == str(k[1]):
                g = k
                sum_flag = 1
        if sum_flag == 1:
            sum_list.append(g[0])
        else:
            sum_list.append(0)
        sum_flag = 0
    # ins_log.read_log('info', sum_list)
    return sum_list


class sqlDelete(BaseHandler):
    def delete(self, *args, **kwargs):
        data = json.loads(self.request.body.decode("utf-8"))
        toid = int(data.get('id', None))  # id号
        with DBContext('w', None, True) as session:
            session.query(AssetSql).filter(AssetSql.id == toid).delete(synchronize_session=False)
            session.commit()
        return self.write(dict(code=0, msg='删除成功'))


class getBarHandler(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        nickname = self.get_current_nickname()
        tostart = self.get_argument('startdate', strip=True) + " " + "00:00:00"  # 要查询的关键字
        toend = self.get_argument('enddate', strip=True) + " " + "23:59:59"  # 要查询的关键字

        with DBContext('r') as session:
            conditions = []
            conditions.append(AssetSql.case_stime >= tostart)
            conditions.append(AssetSql.case_etime <= toend)
            todata = session.query(AssetSql).filter(*conditions).order_by(AssetSql.ctime.desc()).all()
        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["case_num"] = data_dict["case_num"]
            case_dict["case_obj"] = data_dict["case_obj"]
            case_dict["demand_unit"] = data_dict["demand_unit"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_type"] = data_dict["case_type"]
            case_dict["case_ltime"] = data_dict["case_ltime"]
            case_dict["case_name"] = data_dict["case_name"]
            case_dict["case_status"] = data_dict["case_status"]
            case_dict["case_priority"] = data_dict["case_priority"]
            case_dict["demander"] = data_dict["demander"]
            case_dict["case_executor"] = data_dict["case_executor"]
            case_dict["case_source"] = data_dict["case_source"]
            case_dict["case_details"] = data_dict["case_details"]
            case_dict["case_stime"] = str(data_dict["case_stime"])
            case_dict["case_etime"] = str(data_dict["case_etime"])
            case_dict["case_creator"] = data_dict["case_creator"]
            data_list.append(case_dict)

        if len(data_list) > 0:
            laiwen_list = []  # 来问
            upgrade_list = []  # 升级
            fault_list = []  # 故障
            plan_list = []  # 计划工作完成情况
            especially_list = []  # 特急/耗时长的情况
            sudden_list = []  # 其他突发工作
            for i in data_list:
                if i["case_source"] == '来文':
                    laiwen_list.append(i)
                if i["case_type"] == '应用升级':
                    upgrade_list.append(i)
                if i["case_type"] == '故障':
                    fault_list.append(i)
                if i["case_priority"] == '特急' or int(i["case_ltime"]) >= 240:
                    especially_list.append(i)
            other_all = len(data_list) - len(laiwen_list) - len(upgrade_list) - len(fault_list) - len(especially_list)
            bar_list = [
                {"来文": len(laiwen_list), "应用升级": len(upgrade_list), "故障": len(fault_list), "重要工作": len(especially_list),
                 "其他": other_all}]
            pie_list = [{"name": "来文", "value": len(laiwen_list)}, {"name": "应用升级", "value": len(upgrade_list)}, \
                        {"name": "故障", "value": len(fault_list)}, {"name": "重要工作", "value": len(especially_list)},
                        {"name": "其他", "value": other_all}],

            self.write(dict(code=0, msg='获取報告成功', count=1, data=bar_list, list=pie_list))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))

class getSqlIdList(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        with DBContext('r') as session:
            todata = session.query(AssetSql).filter(AssetSql.totype == "定时").order_by(AssetSql.create_time.desc()).all()
            tocount = session.query(AssetSql).filter(AssetSql.totype == "定时").count()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["name"] = data_dict["name"]
            case_dict["header"] = data_dict["header"]
            case_dict["dbname_id"] = data_dict["dbname_id"]
            case_dict["dbname"] = data_dict["dbname"]
            case_dict["totype"] = data_dict["totype"]
            case_dict["sqlstr"] = data_dict["sqlstr"]
            case_dict["remarks"] = data_dict["remarks"]
            case_dict["username"] = data_dict["username"]
            case_dict["obj"] = data_dict["obj"]
            case_dict["department"] = data_dict["department"]
            case_dict["storage"] = data_dict["storage"]
            case_dict["create_time"] = str(data_dict["create_time"])
            data_list.append({"id":case_dict["id"],"name":case_dict["name"]})

        if len(data_list) > 0:
            self.write(dict(code=0, msg='获取成功', count=tocount, data=data_list))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))

class getSqlIdDate(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        value = int(self.get_argument('value', default=None, strip=True))
        with DBContext('r') as session:
            todata = session.query(AssetSql).filter(AssetSql.id == value).all()
            # tocount = session.query(AssetSql).filter().count()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["name"] = data_dict["name"]
            case_dict["header"] = data_dict["header"]
            case_dict["dbname_id"] = data_dict["dbname_id"]
            case_dict["dbname"] = data_dict["dbname"]
            case_dict["totype"] = data_dict["totype"]
            case_dict["sqlstr"] = data_dict["sqlstr"]
            case_dict["remarks"] = data_dict["remarks"]
            case_dict["username"] = data_dict["username"]
            case_dict["obj"] = data_dict["obj"]
            case_dict["department"] = data_dict["department"]
            case_dict["storage"] = data_dict["storage"]
            case_dict["create_time"] = str(data_dict["create_time"])
            data_list.append(case_dict)

        if len(data_list) > 0:
            self.write(dict(code=0, msg='获取成功',  data=data_list))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))

class getdepartmentlist(BaseHandler):
    def get(self, *args, **kwargs):
        objlist = []
        department_list = []
        de_list = []
        obj_list = []
        num = 0
        with DBContext('r') as session:
            todata = session.query(AssetSql).filter().all()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            # case_dict["id"] = data_dict["id"]
            # case_dict["name"] = data_dict["name"]
            # case_dict["header"] = data_dict["header"]
            # case_dict["dbname_id"] = data_dict["dbname_id"]
            # case_dict["dbname"] = data_dict["dbname"]
            # case_dict["totype"] = data_dict["totype"]
            # case_dict["sqlstr"] = data_dict["sqlstr"]
            # case_dict["remarks"] = data_dict["remarks"]
            # case_dict["username"] = data_dict["username"]
            case_dict["obj"] = data_dict["obj"]
            case_dict["department"] = data_dict["department"]
            # case_dict["create_time"] = str(data_dict["create_time"])
            if case_dict["obj"] not in obj_list:
                obj_list.append(case_dict["obj"])
                objlist.append({"k":num,"v":case_dict["obj"]})
            if case_dict["department"] not in de_list:
                de_list.append(case_dict["department"])
                department_list.append({"k":num,"v":case_dict["department"]})
            num  =  num + 1

        if len(department_list) > 0:
            self.write(dict(code=0, msg='获取成功',  data=department_list,objlist=objlist))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[],objlist=[]))


class getstoragelist(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        department = self.get_argument('department', strip=True)
        obj = self.get_argument('obj', strip=True)
        with DBContext('r') as session:
            conditions = []
            conditions.append(AssetSql.department == department)
            conditions.append(AssetSql.obj ==  obj)
            todata = session.query(AssetSql).filter(*conditions).all()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["obj"] = data_dict["obj"]
            case_dict["department"] = data_dict["department"]
            case_dict["storage"] = data_dict["storage"]
            data_list.append({"k":case_dict["id"],"v":case_dict["storage"]})

        if len(data_list) > 0:
            self.write(dict(code=0, msg='获取成功',  data=data_list))
        else:
            self.write(dict(code=-1, msg='没有相关数据', count=0, data=[]))

class getimplementlist(BaseHandler):
    def get(self, *args, **kwargs):
        data_list = []
        data_list2 = []
        start = self.get_argument('start', strip=True)
        end = self.get_argument('end', strip=True)
        storage = self.get_argument('storage', strip=True) #存储过程id
        with DBContext('r') as session:
            conditions = []
            conditions.append(AssetSql.id == int(storage))
            todata = session.query(AssetSql).filter(*conditions).all()

        for msg in todata:
            case_dict = {}
            data_dict = model_to_dict(msg)
            case_dict["id"] = data_dict["id"]
            case_dict["name"] = data_dict["name"]
            case_dict["header"] = data_dict["header"]
            case_dict["dbname_id"] = data_dict["dbname_id"]
            case_dict["dbname"] = data_dict["dbname"]
            case_dict["totype"] = data_dict["totype"]
            case_dict["obj"] = data_dict["obj"]
            case_dict["department"] = data_dict["department"]
            case_dict["storage"] = data_dict["storage"]
            case_dict["create_time"] = str(data_dict["create_time"])
            data_list.append(case_dict)

        if len(data_list) > 0:

            with DBContext('r') as session:
                conditions = []
                conditions.append(DB.id == int(data_list[0]["dbname_id"]))
                DB_data = session.query(DB).filter(*conditions).all()

            for msg in DB_data:
                case_dict = {}
                data_dict = DB_model_to_dict(msg)
                case_dict["id"] = data_dict["id"]
                case_dict["db_instance"] = data_dict["db_instance"] #库名
                case_dict["db_host"] = data_dict["db_host"]
                case_dict["db_port"] = data_dict["db_port"]
                case_dict["db_user"] = data_dict["db_user"]
                case_dict["db_pwd"] = data_dict["db_pwd"]
                case_dict["db_type"] = data_dict["db_type"]  #oracle/mysql

                data_list2.append(case_dict)
            # ins_log.read_log('info', "800000000000000000000000000000000000")
            # ins_log.read_log('info', data_list2)
            # ins_log.read_log('info', "800000000000000000000000000000000000")
            #连接oracle
            #执行存储过程
            #返回数据
            title_list = data_list[0]["header"].split('|')
            columns_list = []
            num = 0
            key_list = []
            for  h in title_list:
                tempstr = 'number'+str(num)
                columns_list.append({ "title": h, "key": tempstr, "editable": "true" })
                key_list.append(tempstr)
                num +=1

            data_list3 =[{"number0":1,"number1":5,"number2":8,"number3":10}]
            ins_log.read_log('info', "800000000000000000000000000000000000")
            ins_log.read_log('info', title_list)
            ins_log.read_log('info', columns_list)
            ins_log.read_log('info', key_list)
            ins_log.read_log('info', "800000000000000000000000000000000000")
            self.write(dict(code=0, msg='获取成功',  data=data_list3,columnslist = columns_list,titlelist=title_list,keylist= key_list))
        else:
            self.write(dict(code=-1, msg='没有相关存储过程数据',data=[],columns_list=[],titlelist=[],keylist=[]))




assetSql_urls = [
    (r"/v1/sql/add/", SqlListHandler),
    (r"/v1/sql/list/", getSqlListHandler),
    (r"/v1/sql/Idlist/", getSqlIdList),
    (r"/v1/sql/IdDate/", getSqlIdDate),
    (r"/v1/sql/getfile/", getCasefileHandler),
    (r"/v1/sql/delete/", sqlDelete),
    (r"/v1/sql/getbar/", getBarHandler),
    (r"/v1/sql/departmentlist/", getdepartmentlist),
    (r"/v1/sql/storagelist/", getstoragelist),
    (r"/v1/sql/implement/", getimplementlist),
]

if __name__ == "__main__":
    pass
